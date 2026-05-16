
import google.generativeai as genaifrom telegram import Update
from telegram.ext import Application, MessageHandler, filters, ContextTypes
from docx import Document
import asyncio

TOKEN = "8750260288:AAGjHjCbD98_s9Z8uZi3TkP7DBY8nrmvPCk"
GEMINI_API = "API_KEY"

genai.configure(api_key=GEMINI_API)

model = genai.GenerativeModel("gemini-1.5-flash")
users = {}

async def reply(update: Update, context: ContextTypes.DEFAULT_TYPE):

    user_id = update.message.from_user.id
    text = update.message.text
response = model.generate_content(
    f"Foydalanuvchiga batafsil va aqlli javob ber: {text}"
)

await update.message.reply_text(response.text)
return
    if user_id not in users:
        users[user_id] = 0

    users[user_id] += 1

    if users[user_id] > 1:
        await update.message.reply_text("⏳ 4 minut kuting...")
        await asyncio.sleep(240)

    doc = Document()

    doc.add_heading("Mustaqil ish", 0)

    content = f"""
Mavzu: {text}

Kirish:
Bu mavzu haqida ma'lumot.

Asosiy qism:
AI yordamida tayyorlandi.

Xulosa:
Tayyor ish tugadi.
"""

    doc.add_paragraph(content)

    filename = "mustaqil_ish.docx"

    doc.save(filename)

    await update.message.reply_document(
        document=open(filename, "rb")
    )

app = Application.builder().token(TOKEN).build()

app.add_handler(
    MessageHandler(filters.TEXT, reply)
)

print("Bot ishladi...")

app.run_polling()
